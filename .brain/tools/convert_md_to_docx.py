#!/usr/bin/env python3
"""
convert_md_to_docx.py - Markdown to Word Converter for Jarvis

Chuyển đổi file Markdown (.md) sang Word (.docx) với formatting đầy đủ.

DECISION GUIDE FOR AI AGENT:

    Khi nào dùng:
    • User yêu cầu export report/analysis sang Word
    • Cần chia sẻ report cho stakeholders không dùng Markdown
    • Convert REPORT.md sau khi phân tích xong

    Cách dùng:
    ┌──────────────────────────────────────────────────────────────────┐
    │ --input report.md                 │ File .md cần convert        │
    │ --output report.docx              │ File .docx output (optional)│
    └──────────────────────────────────────────────────────────────────┘

    Nếu không có --output, tự đổi đuôi .md → .docx

Sử dụng:
    python3 convert_md_to_docx.py --input report.md --output report.docx
    python3 convert_md_to_docx.py --input report.md  # → report.docx

Environment:
    Yêu cầu: venv tại .venv/ (chạy setup.sh hoặc setup.bat để tạo)
    Packages: python-docx, markdown, beautifulsoup4
"""

import argparse
import sys
import os
import subprocess
import tempfile

def _find_venv_python():
    """Tìm Python trong .venv/, hỗ trợ macOS/Linux/Windows."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    kit_root = os.path.dirname(os.path.dirname(script_dir))
    venv_dir = os.path.join(kit_root, ".venv")

    candidates = [
        os.path.join(venv_dir, "bin", "python"),
        os.path.join(venv_dir, "bin", "python3"),
        os.path.join(venv_dir, "Scripts", "python.exe"),
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None

PYTHON_PATH = _find_venv_python()


def convert(input_path: str, output_path: str):
    """Convert .md file to .docx via executor script."""

    if not PYTHON_PATH:
        print("ERROR: Không tìm thấy .venv/. Chạy setup.sh (macOS/Linux) hoặc setup.bat (Windows) trước.")
        sys.exit(1)

    # Kiểm tra input file
    if not os.path.exists(input_path):
        print(f"ERROR: File không tồn tại: {input_path}")
        sys.exit(1)

    # Đọc nội dung MD
    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Lưu MD content vào temp file (tránh escape issues)
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
        f.write(md_content)
        temp_md = f.name

    # Tạo executor script
    executor_script = f'''
import warnings
import sys

warnings.filterwarnings('ignore')

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Helper functions ---

def set_table_borders(table):
    """Set borders cho tất cả cells trong table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing borders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{{edge}}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    tblPr.append(borders)


def set_cell_borders(cell):
    """Set border cho từng cell riêng lẻ (backup)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)

    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{{edge}}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    tcPr.append(borders)


def add_inline_formatting(paragraph, element):
    """Xử lý inline formatting: bold, italic, code, links."""
    if isinstance(element, str):
        # Plain text
        paragraph.add_run(element)
        return

    tag = element.name if hasattr(element, 'name') else None

    if tag is None:
        # NavigableString
        paragraph.add_run(str(element))
        return

    text = element.get_text()

    if tag == 'strong' or tag == 'b':
        run = paragraph.add_run(text)
        run.bold = True
    elif tag == 'em' or tag == 'i':
        run = paragraph.add_run(text)
        run.italic = True
    elif tag == 'code':
        run = paragraph.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    elif tag == 'a':
        href = element.get('href', '')
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        run.underline = True
        if href and href != text:
            paragraph.add_run(f' ({{href}})')
    elif tag == 'br':
        paragraph.add_run('\\n')
    else:
        # Fallback: iterate children
        for child in element.children:
            add_inline_formatting(paragraph, child)


def process_element(doc, element, list_level=0):
    """Xử lý từng HTML element và thêm vào document."""
    if isinstance(element, str):
        text = element.strip()
        if text:
            doc.add_paragraph(text)
        return

    tag = element.name if hasattr(element, 'name') else None
    if tag is None:
        return

    # --- Headings ---
    if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(tag[1])
        heading = doc.add_heading(level=level)
        for child in element.children:
            add_inline_formatting(heading, child)

    # --- Paragraphs ---
    elif tag == 'p':
        para = doc.add_paragraph()
        for child in element.children:
            add_inline_formatting(para, child)

    # --- Code blocks ---
    elif tag == 'pre':
        code_el = element.find('code')
        code_text = code_el.get_text() if code_el else element.get_text()
        para = doc.add_paragraph()
        run = para.add_run(code_text.rstrip())
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        # Set paragraph shading (light gray background)
        pPr = para._p.get_or_add_pPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shading)

    # --- Tables ---
    elif tag == 'table':
        rows_data = []
        # Extract header
        thead = element.find('thead')
        if thead:
            for tr in thead.find_all('tr'):
                cells = [th.get_text().strip() for th in tr.find_all(['th', 'td'])]
                rows_data.append(cells)
        # Extract body
        tbody = element.find('tbody')
        body_rows = tbody.find_all('tr') if tbody else element.find_all('tr')
        if thead:
            # tbody rows only
            for tr in (tbody.find_all('tr') if tbody else []):
                cells = [td.get_text().strip() for td in tr.find_all(['td', 'th'])]
                rows_data.append(cells)
        else:
            for tr in body_rows:
                cells = [td.get_text().strip() for td in tr.find_all(['td', 'th'])]
                rows_data.append(cells)

        if not rows_data:
            return

        # Xác định số cột
        num_cols = max(len(row) for row in rows_data)
        num_rows = len(rows_data)

        # Tạo table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fill data
        for i, row_data in enumerate(rows_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    cell.text = cell_text
                    # Bold header row
                    if i == 0 and thead:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        # Set borders cho table
        set_table_borders(table)

        # Set borders cho từng cell (đảm bảo chắc chắn)
        for row in table.rows:
            for cell in row.cells:
                set_cell_borders(cell)

    # --- Unordered list ---
    elif tag == 'ul':
        for li in element.find_all('li', recursive=False):
            para = doc.add_paragraph(style='List Bullet')
            for child in li.children:
                if hasattr(child, 'name') and child.name in ('ul', 'ol'):
                    # Nested list - simplified: add as indented text
                    for sub_li in child.find_all('li', recursive=False):
                        sub_para = doc.add_paragraph(style='List Bullet 2')
                        sub_para.add_run(sub_li.get_text().strip())
                else:
                    add_inline_formatting(para, child)

    # --- Ordered list ---
    elif tag == 'ol':
        for li in element.find_all('li', recursive=False):
            para = doc.add_paragraph(style='List Number')
            for child in li.children:
                if hasattr(child, 'name') and child.name in ('ul', 'ol'):
                    for sub_li in child.find_all('li', recursive=False):
                        sub_para = doc.add_paragraph(style='List Number 2')
                        sub_para.add_run(sub_li.get_text().strip())
                else:
                    add_inline_formatting(para, child)

    # --- Horizontal rule ---
    elif tag == 'hr':
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('─' * 50)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # --- Blockquote ---
    elif tag == 'blockquote':
        text = element.get_text().strip()
        para = doc.add_paragraph()
        para.style = doc.styles['Intense Quote'] if 'Intense Quote' in [s.name for s in doc.styles] else doc.styles['Normal']
        run = para.add_run(text)
        run.italic = True

    # --- Div or other containers ---
    elif tag in ('div', 'section', 'article', 'main'):
        for child in element.children:
            process_element(doc, child)

    # --- Fallback for unknown tags ---
    else:
        text = element.get_text().strip()
        if text:
            para = doc.add_paragraph()
            para.add_run(text)


# --- Main conversion ---
try:
    # Read markdown
    with open("{temp_md}", "r", encoding="utf-8") as f:
        md_text = f.read()

    # Convert MD → HTML
    html = markdown.markdown(
        md_text,
        extensions=['tables', 'fenced_code', 'codehilite', 'toc']
    )

    # Parse HTML
    soup = BeautifulSoup(html, 'html.parser')

    # Create DOCX
    doc = Document()

    # Set default font to Verdana
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Verdana'
    font.size = Pt(11)

    # Set Verdana for heading styles too
    for i in range(1, 7):
        heading_style_name = f'Heading {{i}}'
        if heading_style_name in [s.name for s in doc.styles]:
            doc.styles[heading_style_name].font.name = 'Verdana'

    # Process all top-level elements
    for element in soup.children:
        process_element(doc, element)

    # Save
    output_path = "{output_path}"
    doc.save(output_path)
    print(f"OK: {{output_path}}")

except Exception as e:
    print(f"ERROR: {{e}}", file=sys.stderr)
    import traceback
    traceback.print_exc(file=sys.stderr)
    sys.exit(1)
'''

    # Lưu script vào temp file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False, encoding='utf-8') as f:
        f.write(executor_script)
        script_file = f.name

    # Chạy script
    try:
        result = subprocess.run(
            [PYTHON_PATH, script_file],
            capture_output=True,
            text=True
        )

        if result.stdout:
            print(result.stdout, end='')
        if result.stderr:
            print(result.stderr, file=sys.stderr, end='')

        return result.returncode

    finally:
        # Cleanup temp files
        try:
            os.unlink(script_file)
            os.unlink(temp_md)
        except:
            pass


def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown to Word (.docx)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:

1. Convert report sang Word:
   python3 convert_md_to_docx.py --input analyses/20260227_analysis/REPORT.md --output report.docx

2. Auto-name output (REPORT.md → REPORT.docx):
   python3 convert_md_to_docx.py --input analyses/20260227_analysis/REPORT.md

Features:
   - Headings (h1-h6) → Word Heading styles
   - **Bold**, *Italic*, `code` inline
   - Tables với BORDER đầy đủ cho tất cả cells
   - Code blocks với font monospace + background
   - Bullet lists, numbered lists
   - Horizontal rules, blockquotes
        """
    )

    parser.add_argument("--input", "-i", required=True, help="Path đến file .md")
    parser.add_argument("--output", "-o", help="Path file .docx output (default: đổi .md → .docx)")

    args = parser.parse_args()

    # Validate paths không chứa ký tự gây lỗi injection
    for path_val, path_name in [(args.input, "--input"), (args.output, "--output")]:
        if path_val and ('"' in path_val or '\\' in path_val):
            parser.error(f'{path_name} không được chứa ký tự " hoặc \\\\')

    # Auto output path
    if not args.output:
        base = os.path.splitext(args.input)[0]
        args.output = base + '.docx'

    return convert(args.input, args.output)


if __name__ == "__main__":
    sys.exit(main())
